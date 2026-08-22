#!/usr/bin/env python3
"""
Odyssey-10 Pro -- Markdown to Word converter.

Purpose-built for this repository's engineering specification. It is not a general
Markdown implementation; it handles exactly the constructs the specification uses, and
it handles them properly:

  * ATX headings (# .. ####), mapped onto Word's built-in heading styles so the
    generated table of contents and the navigation pane both work
  * Pipe tables, rendered as real Word tables with a repeating header row
  * Fenced code blocks and indented ASCII diagrams, in a monospaced style with a light
    background -- the specification is full of block diagrams whose alignment must
    survive the conversion
  * Inline bold, italic and `code`
  * Bulleted and numbered lists, including nesting
  * Block quotes, used in the specification for the safety callouts

Output is BYTE-REPRODUCIBLE: the same Markdown always produces the same .docx. That is
not free -- a .docx is a zip, and python-docx stamps every entry with the current time,
so regenerating an unchanged document used to produce a different file every run.

It mattered for two reasons. Committed binaries churned on every regeneration, so the
history could not distinguish "the specification changed" from "someone re-ran the
converter". And nothing could check whether a Word file was actually up to date, which
is why they twice drifted days behind their Markdown without anyone noticing.

With the timestamps normalised, a differing .docx means the source genuinely changed,
and `check_consistency.py` can regenerate into a temporary file and compare.

Usage:
    python md2docx.py input.md output.docx
    python md2docx.py input.md output.docx --title "..." --subtitle "..."
"""

import argparse
import re
import shutil
import sys
import zipfile
from pathlib import Path

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    raise SystemExit("python-docx is required:  pip install python-docx")


# =====================================================================================
#  Styling
# =====================================================================================
MONO_FONT = "Consolas"
BODY_FONT = "Calibri"
CODE_SHADING = "F2F2F2"
HEADER_SHADING = "1F3864"
ACCENT = RGBColor(0x1F, 0x38, 0x64)


def shade_cell(cell, hex_fill):
    """python-docx has no shading API, so drop down to the underlying XML."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def shade_paragraph(paragraph, hex_fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    p_pr.append(shd)


def repeat_header_row(row):
    """Mark a table row as a header so Word repeats it across page breaks."""
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def build_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for level, size in ((1, 20), (2, 15), (3, 12.5), (4, 11)):
        st = doc.styles[f"Heading {level}"]
        st.font.name = BODY_FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = ACCENT
        st.paragraph_format.space_before = Pt(14 if level <= 2 else 10)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.keep_with_next = True

    code = doc.styles.add_style("OdysseyCode", 1)  # WD_STYLE_TYPE.PARAGRAPH
    code.font.name = MONO_FONT
    code.font.size = Pt(7.5)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.line_spacing = 1.0
    code.paragraph_format.left_indent = Inches(0.15)
    return code


# =====================================================================================
#  Inline formatting
# =====================================================================================
INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*]+?\*)")


def add_inline(paragraph, text):
    """Splits a line into runs, honouring **bold**, *italic* and `code`."""
    # Strip the LaTeX-ish math delimiters the original document used; Word has no
    # equivalent and the escaped forms render as literal backslashes.
    text = text.replace("\\times", "x").replace("\\approx", "~")
    text = re.sub(r"\$\$?(.+?)\$\$?", r"\1", text)
    text = re.sub(r"\\text\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\(pm|le|ge|rightarrow|to)\b",
                  lambda m: {"pm": "+/-", "le": "<=", "ge": ">=",
                             "rightarrow": "->", "to": "->"}[m.group(1)], text)

    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = MONO_FONT
            run.font.size = Pt(9.5)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)


# =====================================================================================
#  Block parsing
# =====================================================================================
def is_table_row(line):
    return line.strip().startswith("|") and line.strip().endswith("|")


def is_table_divider(line):
    s = line.strip()
    return bool(s) and set(s) <= set("|-: ") and "-" in s


def split_row(line):
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def emit_code_block(doc, code_style, lines):
    # Trim leading/trailing blank lines but keep interior alignment exactly.
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()
    if not lines:
        return
    # Remove the common leading indentation so diagrams sit at the left margin.
    indents = [len(l) - len(l.lstrip(" ")) for l in lines if l.strip()]
    strip = min(indents) if indents else 0
    for line in lines:
        p = doc.add_paragraph(style=code_style)
        run = p.add_run(line[strip:] if len(line) > strip else "")
        run.font.name = MONO_FONT
        run.font.size = Pt(7.5)
        shade_paragraph(p, CODE_SHADING)


def emit_table(doc, rows):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j in range(ncols):
            text = row[j] if j < len(row) else ""
            para = cells[j].paragraphs[0]
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.space_before = Pt(2)
            if i == 0:
                run = para.add_run(re.sub(r"\*\*", "", text))
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                shade_cell(cells[j], HEADER_SHADING)
            else:
                add_inline(para, text)
                for r in para.runs:
                    r.font.size = Pt(9)
        if i == 0:
            repeat_header_row(table.rows[0])
    doc.add_paragraph()


def convert(md_text, doc, code_style, heading_shift=0):
    """
    heading_shift lifts every heading by N levels. When the document's leading `#`
    line has been consumed as the title page, its `##` sections are semantically
    top-level, so a shift of 1 puts them on Heading 1 where Word's navigation pane and
    the generated table of contents expect them.
    """
    lines = md_text.split("\n")
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # ---- Fenced code block ------------------------------------------------------
        if stripped.startswith("```"):
            i += 1
            block = []
            while i < n and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1
            emit_code_block(doc, code_style, block)
            continue

        # ---- Indented block (ASCII diagrams) ----------------------------------------
        # Four spaces of indent, and not part of a list.
        if line.startswith("    ") and stripped and not re.match(r"^\s*[-*+]\s", line):
            block = []
            while i < n and (lines[i].startswith("    ") or not lines[i].strip()):
                block.append(lines[i])
                i += 1
                # Stop at two consecutive blanks so we do not swallow the next section.
                if len(block) >= 2 and not block[-1].strip() and not block[-2].strip():
                    break
            emit_code_block(doc, code_style, block)
            continue

        # ---- Table -------------------------------------------------------------------
        if is_table_row(line) and i + 1 < n and is_table_divider(lines[i + 1]):
            rows = [split_row(line)]
            i += 2
            while i < n and is_table_row(lines[i]):
                rows.append(split_row(lines[i]))
                i += 1
            emit_table(doc, rows)
            continue

        # ---- Headings ----------------------------------------------------------------
        m = re.match(r"^(#{1,5})\s+(.*)$", stripped)
        if m:
            level = max(1, min(4, len(m.group(1)) - heading_shift))
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline(p, m.group(2))
            i += 1
            continue

        # Setext-style headings (the original document used underlined headings)
        if (i + 1 < n and stripped
                and re.fullmatch(r"[=]{3,}", lines[i + 1].strip())):
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, stripped)
            i += 2
            continue
        if (i + 1 < n and stripped
                and re.fullmatch(r"[-]{3,}", lines[i + 1].strip())
                and not is_table_divider(lines[i + 1])):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, stripped)
            i += 2
            continue

        # ---- Horizontal rule ----------------------------------------------------------
        if re.fullmatch(r"(\*\s*){3,}|(-\s*){3,}|(_\s*){3,}", stripped):
            doc.add_paragraph("_" * 78).alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # ---- Block quote ---------------------------------------------------------------
        if stripped.startswith(">"):
            block = []
            while i < n and lines[i].strip().startswith(">"):
                block.append(lines[i].strip().lstrip(">").strip())
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_inline(p, " ".join(block))
            for r in p.runs:
                r.italic = True
            shade_paragraph(p, "FFF2CC")
            continue

        # ---- Lists ------------------------------------------------------------------
        m = re.match(r"^(\s*)([-*+])\s+(.*)$", line)
        if m:
            indent = len(m.group(1))
            style = "List Bullet" if indent < 2 else "List Bullet 2"
            p = doc.add_paragraph(style=style)
            add_inline(p, m.group(3))
            i += 1
            continue

        m = re.match(r"^(\s*)\d+[.)]\s+(.*)$", line)
        if m:
            indent = len(m.group(1))
            style = "List Number" if indent < 2 else "List Number 2"
            p = doc.add_paragraph(style=style)
            add_inline(p, m.group(2))
            i += 1
            continue

        # ---- Blank / paragraph ---------------------------------------------------------
        if not stripped:
            i += 1
            continue

        # Collect a paragraph's worth of consecutive non-blank, non-special lines.
        para_lines = []
        while i < n:
            cur = lines[i]
            s = cur.strip()
            if (not s or s.startswith("```") or s.startswith(">")
                    or re.match(r"^#{1,4}\s", s) or is_table_row(cur)
                    or cur.startswith("    ")
                    or re.match(r"^\s*([-*+]|\d+[.)])\s", cur)):
                break
            if (i + 1 < n and re.fullmatch(r"[=\-]{3,}", lines[i + 1].strip())):
                break
            para_lines.append(s)
            i += 1
        if para_lines:
            p = doc.add_paragraph()
            add_inline(p, " ".join(para_lines))


# =====================================================================================
#  Front matter
# =====================================================================================
def add_title_page(doc, title, subtitle, tagline):
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(30)
    r.font.color.rgb = ACCENT
    r.font.name = BODY_FONT

    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(subtitle)
        r.font.size = Pt(14)
        r.italic = True

    doc.add_paragraph()
    if tagline:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(tagline)
        r.font.size = Pt(10.5)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_toc_placeholder(doc):
    """
    Inserts a real Word TOC field. Word populates it on open (or on F9); it cannot be
    pre-rendered from outside Word, so the placeholder text tells the reader what to do.
    """
    doc.add_paragraph("Table of Contents", style="Heading 1")
    p = doc.add_paragraph()
    run = p.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click here and choose 'Update Field' to build the contents."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    for el in (fld_begin, instr, fld_sep, placeholder, fld_end):
        run._r.append(el)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_page_numbers(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for kind, text in (("begin", None), (None, "PAGE"), ("end", None)):
        if kind:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), kind)
        else:
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve")
            el.text = text
        run._r.append(el)


# =====================================================================================
#  The epoch a zip can represent. Any fixed value works; this is the conventional one
#  and is what other reproducible-build tooling uses.
_ZIP_EPOCH = (1980, 1, 1, 0, 0, 0)


def make_reproducible(path):
    """
    Rewrites a .docx so its zip entries carry a fixed timestamp.

    The document's own metadata is already fixed -- python-docx inherits the created and
    modified dates from its template -- so the entry timestamps are the whole of the
    non-determinism. Entry order, names, contents and compression are preserved exactly;
    only the dates change.
    """
    path = Path(path)
    with zipfile.ZipFile(path) as src:
        entries = [(info, src.read(info.filename)) for info in src.infolist()]

    tmp = path.with_suffix(path.suffix + ".tmp")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as out:
        for info, blob in entries:
            fixed = zipfile.ZipInfo(info.filename, date_time=_ZIP_EPOCH)
            fixed.compress_type = info.compress_type
            fixed.external_attr = info.external_attr
            fixed.internal_attr = info.internal_attr
            fixed.create_system = 0          # not the host OS, which also varies
            out.writestr(fixed, blob)

    shutil.move(str(tmp), str(path))


def main():
    ap = argparse.ArgumentParser(description="Convert the Odyssey specification to .docx")
    ap.add_argument("input", type=Path)
    ap.add_argument("output", type=Path)
    ap.add_argument("--title", default=None)
    ap.add_argument("--subtitle", default=None)
    ap.add_argument("--tagline", default="")
    ap.add_argument("--no-toc", action="store_true")
    args = ap.parse_args()

    md = args.input.read_text(encoding="utf-8")

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    code_style = build_styles(doc)

    title = args.title
    heading_shift = 0
    if title is None:
        first = md.split("\n", 1)[0].strip()
        title = re.sub(r"^#+\s*", "", first) or args.input.stem
        # Drop the title line so it is not repeated in the body.
        md = md.split("\n", 1)[1] if "\n" in md else ""
        md = re.sub(r"^={3,}\s*\n", "", md)
        # The `#` line became the title page, so the document's `##` sections are
        # semantically top-level. Lift everything one level.
        heading_shift = 1

    add_title_page(doc, title, args.subtitle, args.tagline)
    if not args.no_toc:
        add_toc_placeholder(doc)
    add_page_numbers(doc)

    convert(md, doc, code_style, heading_shift)

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)
    # Same Markdown in, same bytes out. See make_reproducible().
    make_reproducible(args.output)

    paragraphs = len(doc.paragraphs)
    tables = len(doc.tables)
    print(f"wrote {args.output}  ({paragraphs} paragraphs, {tables} tables)")


if __name__ == "__main__":
    sys.exit(main())
